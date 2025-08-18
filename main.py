import os
import random
from pathlib import Path
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFon
from openai import OpenAI

# =============================
# Check API Key
# =============================
api_key = os.getenv("OPENAI_API_KEY")
if api_key:
    print("✅ LLM mode active (using OpenAI API key)")
    client = OpenAI(api_key=api_key)
else:
    print("⚠ Fallback mode active (no API key found) — generating varied sample questions")
    client = None

# Output folders
output_dir = Path("output")
images_dir = output_dir / "images"
output_dir.mkdir(exist_ok=True)
images_dir.mkdir(exist_ok=True)

# =============================
# Fallback question generator
# =============================
def generate_fallback_questions():
    backpack_colors = random.sample(["Blue", "Green", "Gray", "Red", "Yellow"], 3)
    bottle_types = random.sample(["Stainless", "Plastic", "Glass", "Insulated", "Copper"], 4)
    num_backpacks = len(backpack_colors)
    num_bottles = len(bottle_types)
    total_combinations = num_backpacks * num_bottles

    q1 = f"""@title Combinatorics — Campus Gear Choices
@description Count combinations of independent choices (backpack × bottle).

@question At Greenfield High each student selects **one backpack** and **one water bottle**. 
The backpack colors available are {", ".join(backpack_colors)}. 
The water bottle types available are {", ".join(bottle_types)}. 
How many different backpack–bottle combinations are possible?

@instruction Select the single best answer.
@difficulty easy
@Order 1

@option {total_combinations - 5}
@option {total_combinations - 3}
@option {total_combinations - 1}
@option {total_combinations + 2}
@@option {total_combinations}
@explanation Each combination is formed by choosing one backpack ({num_backpacks} choices) 
and one bottle ({num_bottles} choices). Total combinations = {num_backpacks}×{num_bottles}={total_combinations}.

@subject Quantitative Math
@unit Data Analysis & Probability
@topic Counting & Arrangement Problems
@plusmarks 1"""

    radius = random.choice([2, 3, 4])
    rows, cols = 2, 4
    diameter = radius * 2
    height = diameter
    width = cols * diameter
    length = cols * diameter * 3  # just to vary

    q2 = f"""@title Packed Spheres — Box Dimensions
@description Determine the dimensions of a rectangular box tightly packed with identical spheres arranged in a rectangular grid (top view).

@question A rectangular box contains {rows*cols} identical spheres arranged in {rows} rows and {cols} columns (each sphere touches its neighbors). Each sphere has radius \\({radius}\\) centimeters. Which of the following is closest to the internal dimensions (in centimeters) of the rectangular box (height × width × length)?

@instruction Choose the option that lists the box dimensions in centimeters.
@difficulty moderate
@Order 2

@option \\({height} \\times {width} \\times {length - 5}\\)
@option \\({height + 2} \\times {width} \\times {length}\\)
@@option \\({height} \\times {width} \\times {length}\\)
@option \\({height + 4} \\times {width + 6} \\times {length + 12}\\)
@option \\({height + 3} \\times {width} \\times {length + 6}\\)
@explanation Each sphere has diameter \\(d=2r=2\\times{radius}={diameter}\\) cm. For a {rows}-by-{cols} tight grid: Height={height} cm, Width={width} cm, Length={length} cm. So dimensions = \\({height} \\times {width} \\times {length}\\).

@subject Quantitative Math
@unit Geometry and Measurement
@topic Solid Figures (Volume of Cubes)
@plusmarks 1"""

    return q1, q2

# =============================
# Real LLM question generator
# =============================
def generate_llm_questions():
    prompt = """Generate 2 unique math questions in the following strict format:

1. Each question should start with:
@title ...
@description ...
@question ...
@instruction ...
@difficulty ...
@Order ...
@option ...
@@option ... (the correct one)
@explanation ...
@subject ...
@unit ...
@topic ...
@plusmarks ...

2. The first question should be about combinatorics or probability.
3. The second question should be about geometry and include dimensions or volume.
4. Keep difficulty between easy and moderate.
"""
    response = client.chat.completions.create(
        model="gpt-4o-mini",  # You can change model
        messages=[
            {"role": "system", "content": "You are an assistant that generates math questions in a strict tagging format."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.7
    )

    content = response.choices[0].message.content
    # Split into two questions using tags
    questions = content.strip().split("\n\n")
    if len(questions) >= 2:
        return questions[0], questions[1]
    else:
        return content, ""

# =============================
# Create diagram
# =============================
def create_diagram():
    img = Image.new("RGB", (800, 300), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("DejaVuSans.ttf", 18)
    except:
        font = ImageFont.load_default()

    sphere_radius_px = 40
    padding_px = 20
    start_x = padding_px
    start_y = padding_px

    for row in range(2):
        for col in range(4):
            cx = start_x + col * (sphere_radius_px * 2 + padding_px) + sphere_radius_px
            cy = start_y + row * (sphere_radius_px * 2 + padding_px) + sphere_radius_px
            draw.ellipse(
                [cx - sphere_radius_px, cy - sphere_radius_px,
                 cx + sphere_radius_px, cy + sphere_radius_px],
                fill=(173, 216, 230), outline=(0, 0, 0)
            )

    draw.text((20, 250), "Top view: 2 rows × 4 columns of spheres", fill=(0, 0, 0), font=font)

    img_path = images_dir / "packed_spheres_diagram.png"
    img.save(img_path)
    return img_path

# =============================
# Main execution
# =============================
if api_key:
    try:
        q1, q2 = generate_llm_questions()
    except Exception as e:
        print(f"⚠ LLM generation failed: {e}")
        q1, q2 = generate_fallback_questions()
else:
    q1, q2 = generate_fallback_questions()

# Create document
doc = Document()
doc.add_heading("Math Question Generation — Output", level=1)
doc.add_paragraph(q1)
doc.add_page_break()
doc.add_paragraph(q2)
doc.add_paragraph("\nDiagram:")
doc.add_picture(str(create_diagram()), width=Inches(4))

output_docx = output_dir / "questions_output_with_diagram.docx"
doc.save(output_docx)

print(f"📄 Saved generated questions to: {output_docx}")

